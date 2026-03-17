# ==========================================
# 工具函数：数据处理、CSS 加载等
# ==========================================

import pandas as pd
from pathlib import Path
import re
from .config import EMOTION_NAMES, EMOTION_STYLES, COARSE_MAP


def load_css(css_path: str) -> str:
    """
    加载 CSS 文件内容
    
    Args:
        css_path: CSS 文件路径
        
    Returns:
        CSS 内容字符串
    """
    path = Path(css_path)
    if path.exists():
        return path.read_text(encoding='utf-8')
    return ""


def load_csv_with_encoding(file_path, encodings=None):
    """
    尝试多种编码读取 CSV 文件
    
    Args:
        file_path: 文件路径或文件对象
        encodings: 编码列表，默认使用常见中文编码
        
    Returns:
        DataFrame 或 None
    """
    if encodings is None:
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "big5", "cp936", "latin1"]
    
    last_error = None
    
    for enc in encodings:
        try:
            if hasattr(file_path, 'seek'):
                # 文件对象需要重置指针
                file_path.seek(0)
                df = pd.read_csv(file_path, encoding=enc)
            else:
                df = pd.read_csv(file_path, encoding=enc)
            return df, None
        except Exception as e:
            last_error = e
            continue
    
    return None, last_error


def validate_dataframe(df, required_columns=None):
    """
    验证 DataFrame 是否包含必需的列
    
    Args:
        df: DataFrame
        required_columns: 必需列名列表，默认为 ['text']
        
    Returns:
        (bool, str): 是否有效和错误信息
    """
    if required_columns is None:
        required_columns = ['text']
    
    if df is None:
        return False, "DataFrame 为空"
    
    missing_cols = [col for col in required_columns if col not in df.columns]
    if missing_cols:
        return False, f"缺少必需的列：{', '.join(missing_cols)}"
    
    return True, ""


def format_batch_results(results: list) -> pd.DataFrame:
    """
    格式化批量分析结果
    
    Args:
        results: 预测结果列表
        
    Returns:
        DataFrame
    """
    formatted = []
    for r in results:
        row = {
            "text": r.get("text", ""),
            "fine_emotion": r.get("fine_name") or r.get("fine_emotion", ""),
            "coarse_emotion": r.get("coarse") or r.get("coarse_emotion", ""),
            "confidence": r.get("confidence", 0.0),
            **{f"prob_{k}": v for k, v in r.get("all_probabilities", {}).items()},
        }
        formatted.append(row)
    
    return pd.DataFrame(formatted)


def get_example_texts() -> list:
    """
    获取示例文本列表
    
    Returns:
        示例文本列表
    """
    return [
        ("开心示例", "今天终于拿到了心仪的 offer，感觉所有的努力都值得了！"),
        ("悲伤示例", "听到这个坏消息，我心里非常难过。"),
        ("惊讶示例", "居然中了大奖，太意外了！"),
    ]


def emotion_id_by_name(emotion_name: str) -> int:
    """
    根据情感名称获取情感 ID
    
    Args:
        emotion_name: 情感名称
        
    Returns:
        情感 ID，找不到返回 0
    """
    for k, v in EMOTION_NAMES.items():
        if v == emotion_name:
            return int(k)
    return 0


def get_coarse_badge_class(coarse: str) -> str:
    """
    获取粗粒度情感的 CSS 类名
    
    Args:
        coarse: 粗粒度情感标签（正面/负面/中性）
        
    Returns:
        CSS 类名
    """
    if coarse == "正面":
        return "positive"
    if coarse == "负面":
        return "negative"
    return "neutral"


def split_chinese_sentences(text: str) -> list:
    """
    
    支持多种中文标点符号（句号、感叹号、问号等）进行句子分割
    
    Args:
        text: 输入文本
        
    Returns:
        句子列表
    """
    import re

    # 去除换行符，将文本合并为连续的一行
    text = text.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')
    
    # 中文句子结束标点：。！？!? (包括中英文)
    sentence_end_pattern = r'([.。！？!?])'
    
    sentences = []
    
    # 使用正则表达式分割句子
    parts = re.split(sentence_end_pattern, text)
    
    # 重组句子和标点
    current_sentence = ""
    for i, part in enumerate(parts):
        if not part:
            continue
            
        # 如果是标点符号
        if re.match(sentence_end_pattern, part):
            current_sentence += part
            sentences.append(current_sentence.strip())
            current_sentence = ""
        else:
            current_sentence += part
    
    # 处理最后剩余的部分
    if current_sentence.strip():
        sentences.append(current_sentence.strip())
    
    return sentences


def extract_text_from_pdf(file_path) -> str:
    """
    从 PDF 文件中提取文本
    
    Args:
        file_path: PDF 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    import pdfplumber
    
    try:
        if hasattr(file_path, 'seek'):
            # 文件对象
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        else:
            # 文件路径
            with pdfplumber.open(str(file_path)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"无法读取 PDF 文件：{str(e)}")


def extract_text_from_docx(file_path) -> str:
    """
    从 DOCX 文件中提取文本
    
    Args:
        file_path: DOCX 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    from docx import Document
    
    try:
        if hasattr(file_path, 'seek'):
            # 文件对象需要保存到临时文件
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(file_path.read())
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            import os
            os.unlink(tmp_path)
        else:
            doc = Document(str(file_path))
        
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        raise Exception(f"无法读取 DOCX 文件：{str(e)}")


def extract_text_from_md(file_path) -> str:
    """
    从 Markdown 文件中提取纯文本
    
    Args:
        file_path: MD 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    import markdown
    from bs4 import BeautifulSoup
    
    try:
        if hasattr(file_path, 'seek'):
            # 文件对象
            md_content = file_path.read().decode('utf-8') if hasattr(file_path, 'read') else file_path.read()
            md_text = str(md_content)
        else:
            # 文件路径
            with open(str(file_path), 'r', encoding='utf-8') as f:
                md_text = f.read()
        
        # 将 Markdown 转换为 HTML
        html = markdown.markdown(md_text)
        
        # 从 HTML 中提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        
        # 清理多余空白
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)
    except Exception as e:
        raise Exception(f"无法读取 Markdown 文件：{str(e)}")


def extract_text_from_txt(file_path) -> str:
    """
    从 TXT 文件中提取文本
    
    Args:
        file_path: TXT 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'latin1']
    last_error = None
    
    for encoding in encodings:
        try:
            if hasattr(file_path, 'seek'):
                file_path.seek(0)
                content = file_path.read().decode(encoding) if hasattr(file_path, 'read') else str(file_path.read())
            else:
                with open(str(file_path), 'r', encoding=encoding) as f:
                    content = f.read()
            
            return content.strip()
        except Exception as e:
            last_error = e
            continue
    
    raise Exception(f"无法读取 TXT 文件（尝试编码：{encodings}）：{str(last_error)}")


def process_document_to_sentences(file_content, file_type: str) -> pd.DataFrame:
    """
    处理文档文件为句子 DataFrame
    
    Args:
        file_content: 文件内容（字节或文件对象）
        file_type: 文件类型 ('pdf', 'txt', 'docx', 'md')
        
    Returns:
        包含 'text' 列的 DataFrame
    """
    import tempfile
    import os
    
    # 创建临时文件
    suffix_map = {'pdf': '.pdf', 'txt': '.txt', 'docx': '.docx', 'md': '.md'}
    suffix = suffix_map.get(file_type, '')
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        if isinstance(file_content, bytes):
            tmp.write(file_content)
        else:
            # 假设是文件对象
            file_content.seek(0)
            tmp.write(file_content.read())
        tmp_path = tmp.name
    
    try:
        # 提取文本
        if file_type == 'pdf':
            text = extract_text_from_pdf(tmp_path)
        elif file_type == 'txt':
            text = extract_text_from_txt(tmp_path)
        elif file_type == 'docx':
            text = extract_text_from_docx(tmp_path)
        elif file_type == 'md':
            text = extract_text_from_md(tmp_path)
        else:
            raise ValueError(f"不支持的文件类型：{file_type}")
        
        # 划分句子
        sentences = split_chinese_sentences(text)
        
        # 创建 DataFrame
        df = pd.DataFrame({'text': sentences})
        
        return df
    finally:
        # 清理临时文件
        try:
            os.unlink(tmp_path)
        except:
            pass


def load_excel_with_encoding(file_path, encodings=None):
    """
    读取 Excel 文件（支持多种编码）
    
    Args:
        file_path: 文件路径或文件对象
        encodings: 编码列表（用于读取 CSV 风格的 Excel）
        
    Returns:
        DataFrame 或 None
    """
    try:
        if hasattr(file_path, 'seek'):
            file_path.seek(0)
            df = pd.read_excel(file_path, engine='openpyxl')
        else:
            df = pd.read_excel(str(file_path), engine='openpyxl')
        return df
    except Exception as e:
        return None, e