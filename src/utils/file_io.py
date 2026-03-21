# ==========================================
# 文件 I/O 工具：读取各种格式的文件
# ==========================================

import pandas as pd
from pathlib import Path


def load_css(css_path: str) -> str:
    """
    加载 CSS 文件内容
    
    Args:
        css_path: CSS 文件路径
        
    Returns:
        CSS 内容字符串
    """
    path = Path(css_path)
    if path.exists():
        return path.read_text(encoding='utf-8')
    return ""


def load_csv_with_encoding(file_path, encodings=None):
    """
    尝试多种编码读取 CSV 文件
    
    Args:
        file_path: 文件路径或文件对象
        encodings: 编码列表，默认使用常见中文编码
        
    Returns:
        DataFrame 或 None
    """
    if encodings is None:
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "big5", "cp936", "latin1"]
    
    last_error = None
    
    for enc in encodings:
        try:
            if hasattr(file_path, 'seek'):
                # 文件对象需要重置指针
                file_path.seek(0)
                df = pd.read_csv(file_path, encoding=enc)
            else:
                df = pd.read_csv(file_path, encoding=enc)
            return df, None
        except Exception as e:
            last_error = e
            continue
    
    return None, last_error


def load_excel_with_encoding(file_path, encodings=None):
    """
    读取 Excel 文件（支持多种编码）
    
    Args:
        file_path: 文件路径或文件对象
        encodings: 编码列表（用于读取 CSV 风格的 Excel）
        
    Returns:
        DataFrame 或 None
    """
    try:
        if hasattr(file_path, 'seek'):
            file_path.seek(0)
            df = pd.read_excel(file_path, engine='openpyxl')
        else:
            df = pd.read_excel(str(file_path), engine='openpyxl')
        return df
    except Exception as e:
        return None, e


def extract_text_from_pdf(file_path) -> str:
    """
    从 PDF 文件中提取文本
    
    Args:
        file_path: PDF 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    import pdfplumber
    
    try:
        if hasattr(file_path, 'seek'):
            # 文件对象
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        else:
            # 文件路径
            with pdfplumber.open(str(file_path)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"无法读取 PDF 文件：{str(e)}")


def extract_text_from_docx(file_path) -> str:
    """
    从 DOCX 文件中提取文本
    
    Args:
        file_path: DOCX 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    from docx import Document
    
    try:
        if hasattr(file_path, 'seek'):
            # 文件对象需要保存到临时文件
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(file_path.read())
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            import os
            os.unlink(tmp_path)
        else:
            doc = Document(str(file_path))
        
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        raise Exception(f"无法读取 DOCX 文件：{str(e)}")


def extract_text_from_md(file_path) -> str:
    """
    从 Markdown 文件中提取纯文本
    
    Args:
        file_path: MD 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    import markdown
    from bs4 import BeautifulSoup
    
    try:
        if hasattr(file_path, 'seek'):
            # 文件对象
            md_content = file_path.read().decode('utf-8') if hasattr(file_path, 'read') else file_path.read()
            md_text = str(md_content)
        else:
            # 文件路径
            with open(str(file_path), 'r', encoding='utf-8') as f:
                md_text = f.read()
        
        # 将 Markdown 转换为 HTML
        html = markdown.markdown(md_text)
        
        # 从 HTML 中提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        
        # 清理多余空白
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)
    except Exception as e:
        raise Exception(f"无法读取 Markdown 文件：{str(e)}")


def extract_text_from_txt(file_path) -> str:
    """
    从 TXT 文件中提取文本
    
    Args:
        file_path: TXT 文件路径或文件对象
        
    Returns:
        提取的文本内容
    """
    encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'latin1']
    last_error = None
    
    for encoding in encodings:
        try:
            if hasattr(file_path, 'seek'):
                file_path.seek(0)
                content = file_path.read().decode(encoding) if hasattr(file_path, 'read') else str(file_path.read())
            else:
                with open(str(file_path), 'r', encoding=encoding) as f:
                    content = f.read()
            
            return content.strip()
        except Exception as e:
            last_error = e
            continue
    
    raise Exception(f"无法读取 TXT 文件（尝试编码：{encodings}）：{str(last_error)}")


def read_csv_file(file_path):
    """
    读取 CSV 文件
    
    Args:
        file_path: CSV 文件路径
        
    Returns:
        DataFrame
    """
    df, error = load_csv_with_encoding(file_path)
    if error is not None:
        raise Exception(f"读取 CSV 文件失败：{str(error)}")
    return df


def write_csv_file(df, file_path):
    """
    写入 CSV 文件
    
    Args:
        df: DataFrame
        file_path: 输出文件路径
    """
    df.to_csv(file_path, index=False, encoding='utf-8-sig')
